import os
import json
import logging
from datetime import datetime
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from src.document_creator.docx_generator import DocxResumeGenerator

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Google Drive API scope
SCOPES = ['https://www.googleapis.com/auth/drive.file']

class TestDocxGenerator:
    """Test class for document generation"""
    def __init__(self):
        self.logger = logging.getLogger('TestDocxGenerator')

    def create_test_document(self):
        """Create a test document with sample content matching the desired format"""
        doc = Document()
        
        # Set up the page
        section = doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

        # Create a table for the two-column layout
        table = doc.add_table(rows=1, cols=2)
        table.allow_autofit = False
        
        # Set the table to extend to page edges
        table_pr = table._element.xpath('w:tblPr')[0]
        table_width = OxmlElement('w:tblW')
        table_width.set(qn('w:w'), str(int(section.page_width.twips)))
        table_width.set(qn('w:type'), 'dxa')
        table_pr.append(table_width)
        
        # Set column widths (matching screenshot proportions)
        widths = [2.5, 6.0]  # Left column, Right column
        for idx, width in enumerate(widths):
            tc = table.rows[0].cells[idx]
            tcp = tc._element.tcPr
            tcw = OxmlElement('w:tcW')
            tcw.set(qn('w:w'), str(int(Inches(width).twips)))
            tcw.set(qn('w:type'), 'dxa')
            tcp.append(tcw)
        
        # Get cells
        left_cell = table.rows[0].cells[0]
        right_cell = table.rows[0].cells[1]
        
        # Set background color for left column (light blue-grey)
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E6EAF0"/>')
        left_cell._tc.get_or_add_tcPr().append(shading_elm)
        
        # Add cell margins
        def set_cell_margins(cell, top=0.2, right=0.4, bottom=0.2, left=0.4):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            
            for side, width in [('top', top), ('right', right), ('bottom', bottom), ('left', left)]:
                node = OxmlElement(f'w:{side}')
                node.set(qn('w:w'), str(int(Inches(width).twips)))
                node.set(qn('w:type'), 'dxa')
                tcMar.append(node)
            
            tcPr.append(tcMar)
        
        set_cell_margins(left_cell)
        set_cell_margins(right_cell)
        
        # Add sample content with proper formatting
        # Left column
        name = left_cell.add_paragraph()
        name_run = name.add_run("Test Name")
        name_run.bold = True
        name_run.font.size = Pt(24)
        
        title = left_cell.add_paragraph()
        title_run = title.add_run("Test Title")
        title_run.font.size = Pt(16)
        
        # Right column with grey header
        exp_header = right_cell.add_paragraph()
        exp_header_run = exp_header.add_run("PROFESSIONAL EXPERIENCE")
        exp_header_run.bold = True
        exp_header._p.get_or_add_pPr().append(
            parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
        )
        
        return doc

def load_resume_data():
    """Load resume data from JSON file"""
    logger = logging.getLogger('TestDocxGenerator')
    try:
        with open('assets/resume_data.json', 'r') as f:
            data = json.load(f)
        logger.info("Resume data loaded successfully")
        return data
    except FileNotFoundError:
        logger.error("Resume data file not found at assets/resume_data.json")
        raise
    except json.JSONDecodeError:
        logger.error("Invalid JSON format in resume data file")
        raise
    except Exception as e:
        logger.error(f"Unexpected error loading resume data: {str(e)}")
        raise

def test_document_formatting():
    """Test document formatting with comprehensive error handling"""
    logger = logging.getLogger('TestDocxGenerator')
    print("\nTesting Document Formatting...")
    print("=" * 50)
    
    try:
        # Create test directories if they don't exist
        os.makedirs("generated_resumes", exist_ok=True)
        
        # Initialize test generator
        test_generator = TestDocxGenerator()
        
        # Create a test document
        logger.info("Creating test document with new formatting...")
        doc = test_generator.create_test_document()
        
        # Save test document
        test_path = os.path.join("generated_resumes", "test_formatting.docx")
        try:
            doc.save(test_path)
            logger.info(f"Test document saved successfully at {test_path}")
            print("\nTest document created with:")
            print("✓ Two-column layout (2.5\" : 6.0\")")
            print("✓ Light blue-grey background (E6EAF0) in left column")
            print("✓ Grey section headers (F2F2F2) in right column")
            print("✓ Proper text margins and spacing")
            print(f"✓ Saved at: {test_path}")
        except Exception as e:
            logger.error(f"Failed to save test document: {str(e)}")
            raise
        
        # Initialize resume generator
        generator = DocxResumeGenerator()
        
        # Load resume data
        logger.info("Loading resume data...")
        resume_data = load_resume_data()
        
        # Verify required sections
        required_sections = ["personal_info", "experience"]
        missing_sections = [section for section in required_sections if section not in resume_data]
        if missing_sections:
            logger.warning(f"Missing required sections in resume data: {', '.join(missing_sections)}")
        
        print("\nGenerating full resume with:")
        print("- Exact column proportions (2.5\" : 6.0\")")
        print("- Matched blue-grey background (E6EAF0)")
        print("- Grey section headers (F2F2F2)")
        print("- Proper font sizes (24pt name, 16pt title)")
        print("- Correct margins and spacing")
        
        # Test metadata
        metadata = {
            "job_title": "Python Developer",
            "company": "Test Company",
            "job_id": "TEST123"
        }
        
        # Generate document with real data
        logger.info("Generating resume document with real data...")
        output_path = generator.generate_resume_file(resume_data, metadata)
        
        if output_path and os.path.exists(output_path):
            file_size = os.path.getsize(output_path)
            print(f"\n✓ Generated resume at: {output_path}")
            logger.info(f"Document created successfully: {output_path} ({file_size} bytes)")
            
            # Upload to Google Drive and get sharing link
            print("\nGoogle Drive Integration:")
            if generator.drive_service:
                try:
                    print("Uploading document to Google Drive...")
                    file_id = generator._upload_to_drive(output_path)
                    if file_id:
                        print(f"✓ Document uploaded with ID: {file_id}")
                        print("Creating sharing permissions...")
                        
                        # Create sharing permission
                        permission = {
                            'type': 'anyone',
                            'role': 'reader'
                        }
                        generator.drive_service.permissions().create(
                            fileId=file_id,
                            body=permission
                        ).execute()
                        
                        # Get sharing link
                        file = generator.drive_service.files().get(
                            fileId=file_id,
                            fields='webViewLink'
                        ).execute()
                        
                        print(f"✓ Document shared successfully")
                        print(f"✓ View document at: {file.get('webViewLink')}")
                    else:
                        print("✗ Failed to upload document")
                except Exception as e:
                    logger.error(f"Error with Google Drive integration: {str(e)}")
                    print(f"✗ Error: {str(e)}")
            else:
                print("✗ Google Drive integration not enabled")
                print("  To enable, set GOOGLE_DRIVE_ENABLED=true in your .env file")
            
            print("\nPlease verify the following in the generated document:")
            print("1. Left column width is exactly 2.5 inches")
            print("2. Background colors match the screenshot:")
            print("   - Left column: E6EAF0 (light blue-grey)")
            print("   - Section headers: F2F2F2 (light grey)")
            print("3. Font sizes are correct:")
            print("   - Name: 24pt")
            print("   - Title: 16pt")
            print("4. Margins and spacing match the screenshot")
            
            # Verify document sections
            sections_included = []
            if resume_data.get("personal_info"):
                sections_included.append("Personal Info")
            if resume_data.get("summary"):
                sections_included.append("Summary")
            if resume_data.get("experience"):
                sections_included.append("Experience")
            if resume_data.get("skills"):
                sections_included.append("Skills")
            if resume_data.get("projects"):
                sections_included.append("Projects")
            
            print("\nIncluded sections:")
            for section in sections_included:
                print(f"- {section}")
        else:
            logger.error("Failed to create document or document not found")
            print("✗ Failed to create test document")
            
    except Exception as e:
        logger.error(f"Error during testing: {str(e)}")
        print(f"✗ Error during testing: {str(e)}")
        raise

def test_font_installation():
    """Test font installation and availability"""
    logger = logging.getLogger('TestDocxGenerator')
    print("\nTesting Font Installation...")
    print("=" * 50)
    
    try:
        generator = DocxResumeGenerator()
        font_path = "assets/fonts/ZillaSlab-Medium.ttf"
        
        if os.path.exists(font_path):
            print("✓ Font file found")
            logger.info(f"Font file found at {font_path}")
        else:
            print("✗ Font file not found")
            logger.warning(f"Font file not found at {font_path}")
            print("\nPlease ensure the Zilla Slab Medium font is:")
            print("1. Downloaded from Google Fonts")
            print("2. Placed in assets/fonts/ZillaSlab-Medium.ttf")
            
    except Exception as e:
        logger.error(f"Error testing font installation: {str(e)}")
        print(f"✗ Error testing font installation: {str(e)}")
        raise

def main():
    """Main test function"""
    logger = logging.getLogger('TestDocxGenerator')
    try:
        # Ensure required directories exist
        os.makedirs("assets/fonts", exist_ok=True)
        os.makedirs("generated_resumes", exist_ok=True)
        
        # Run tests
        test_font_installation()
        test_document_formatting()
        
    except Exception as e:
        logger.error(f"Test suite failed: {str(e)}")
        print(f"✗ Test suite failed: {str(e)}")
        raise

if __name__ == "__main__":
    load_dotenv()
    main()