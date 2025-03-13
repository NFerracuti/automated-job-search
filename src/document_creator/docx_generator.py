import os
import logging
import json
from datetime import datetime
from typing import Dict, List, Any

from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
import shutil
import platform
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request

from dotenv import load_dotenv

load_dotenv()

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger('DocxGenerator')

# Google Drive API scope
SCOPES = ['https://www.googleapis.com/auth/drive.file']

class DocxResumeGenerator:
    """Class to generate Word documents from resume data"""
    
    def __init__(self, config_path='config.json', template_path=None):
        # Load config
        with open(config_path, 'r') as f:
            self.config = json.load(f)
        
        # Template path
        self.template_path = template_path or os.getenv('RESUME_TEMPLATE_PATH', 'assets/resume_template.docx')
        
        # Initialize logger
        self.logger = logging.getLogger('DocxGenerator')
        
        # Initialize font
        self._install_font()
        
        # Initialize Google Drive service
        drive_enabled = os.getenv('GOOGLE_DRIVE_ENABLED', 'false').lower() == 'true'
        if drive_enabled:
            self.logger.info("Google Drive integration enabled in .env")
            self.drive_service = self._init_drive_service('google-service-account.json')
            if self.drive_service:
                self.logger.info("Google Drive service initialized successfully")
            else:
                self.logger.warning("Failed to initialize Google Drive service")
        else:
            self.logger.info("Google Drive integration disabled in .env")
            self.drive_service = None
        
        self.logger.info("Docx Resume Generator initialized")
        
        # Store resume data for later use
        try:
            with open('assets/resume_data.json', 'r') as f:
                self.resume_data = json.load(f)
                self.logger.info(f"Loaded resume_data.json with keys: {self.resume_data.keys()}")
        except Exception as e:
            self.logger.error(f"Failed to load resume_data.json: {str(e)}")
            self.resume_data = {}
    
    def _install_font(self):
        """Install the Zilla Slab font if not already installed"""
        try:
            font_path = os.path.join('assets', 'fonts', 'ZillaSlab-Medium.ttf')
            if not os.path.exists(font_path):
                logger.warning(f"Font file not found at {font_path}")
                return

            if platform.system() == "Windows":
                import ctypes
                from ctypes import windll, byref, create_unicode_buffer
                
                FR_PRIVATE = 0x10
                FR_NOT_ENUM = 0x20
                
                font_path_buffer = create_unicode_buffer(os.path.abspath(font_path))
                flags = FR_PRIVATE | FR_NOT_ENUM
                if not windll.gdi32.AddFontResourceExW(byref(font_path_buffer), flags, 0):
                    logger.warning("Failed to install font on Windows")
            
            elif platform.system() == "Darwin":  # macOS
                user_font_dir = os.path.expanduser('~/Library/Fonts')
                if not os.path.exists(os.path.join(user_font_dir, 'ZillaSlab-Medium.ttf')):
                    shutil.copy2(font_path, user_font_dir)
            
            else:  # Linux
                font_dir = os.path.expanduser('~/.local/share/fonts')
                os.makedirs(font_dir, exist_ok=True)
                if not os.path.exists(os.path.join(font_dir, 'ZillaSlab-Medium.ttf')):
                    shutil.copy2(font_path, font_dir)
                    os.system('fc-cache -f -v')
            
            logger.info("Zilla Slab font installed successfully")
            
        except Exception as e:
            logger.warning(f"Failed to install font: {str(e)}")
    
    def _create_basic_resume(self, resume_data):
        """Create a resume document from scratch without using a template"""
        self.logger.info("DIAGNOSTICS - Starting _create_basic_resume")
        self.logger.info(f"DIAGNOSTICS - Skills in _create_basic_resume: <<<{resume_data.get('skills', 'NOT FOUND')}>>> (type: {type(resume_data.get('skills', ''))})")
        
        doc = Document()
        
        # Set up the page with reduced margins
        section = doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        
        # Create a table for the two-column layout
        table = doc.add_table(rows=1, cols=2)
        table.allow_autofit = False
        
        # Set column widths
        for idx, width in enumerate([3.0, 4.5]):
            tc = table.rows[0].cells[idx]
            tcp = tc._element.tcPr
            tcw = OxmlElement('w:tcW')
            tcw.set(qn('w:w'), str(int(Inches(width).twips)))
            tcw.set(qn('w:type'), 'dxa')
            tcp.append(tcw)
        
        # Get cells
        left_cell = table.rows[0].cells[0]
        right_cell = table.rows[0].cells[1]
        
        # Remove any default paragraph in right cell
        if right_cell.paragraphs:
            p = right_cell.paragraphs[0]._element
            p.getparent().remove(p)
        
        # Add vertical line between columns
        table.style = 'Table Grid'
        # Remove all borders except the vertical line between columns
        for cell in table.rows[0].cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border in ['top', 'bottom', 'left', 'right']:
                border_el = OxmlElement(f'w:{border}')
                border_el.set(qn('w:val'), 'nil')
                if (border == 'right' and cell == left_cell) or (border == 'left' and cell == right_cell):
                    border_el.set(qn('w:val'), 'single')
                    border_el.set(qn('w:sz'), '4')  # Thin line
                    border_el.set(qn('w:color'), '000000')
                tcBorders.append(border_el)
            tcPr.append(tcBorders)
        
        # Set custom font for the entire document
        font_name = "Zilla Slab Medium"
        styles = doc.styles
        style = styles['Normal']
        style.font.name = font_name
        style.font.size = Pt(10)  # Set base font size to 10pt
        
        # Function to create highlighted section header
        def add_section_header(cell, text, add_space_before=True):
            # Add blank line before header (optional)
            if add_space_before:
                spacer_before = cell.add_paragraph()
                spacer_before.paragraph_format.space_after = Pt(0)
            
            # Add header with grey background
            header = cell.add_paragraph()
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = header.add_run(text)
            run.bold = True
            # Add light grey background only to the text line
            pPr = header._p.get_or_add_pPr()
            pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>'))
            header.paragraph_format.space_after = Pt(0)  # Remove spacing from header
            
            # Add white spacing below header
            spacer_after = cell.add_paragraph()
            spacer_after.paragraph_format.space_before = Pt(0)
            spacer_after.paragraph_format.space_after = Pt(0)
            
            return header
        
        # Add content to left column
        personal = resume_data.get("personal_info", {})
        
        # Name
        name_paragraph = left_cell.paragraphs[0]
        run = name_paragraph.add_run("Nick Ferracuti")
        run.bold = True
        run.font.size = Pt(18)
        name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Title
        title_paragraph = left_cell.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_paragraph.add_run(personal.get("title", "Full Stack Developer"))
        run.bold = True
        
        # Contact info with hyperlinks
        contact_paragraph = left_cell.add_paragraph()
        contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_paragraph.paragraph_format.space_after = Pt(0)
        contact_paragraph.paragraph_format.space_before = Pt(0)
        
        # Website with hyperlink
        website_run = contact_paragraph.add_run()
        website_run.text = "www.nickferracuti.com"
        website_run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color for hyperlink
        website_run.font.underline = True
        # Add hyperlink relationship
        rel_id = doc.part.relate_to('http://www.nickferracuti.com', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), rel_id)
        hyperlink.append(website_run._r)
        contact_paragraph._p.append(hyperlink)
        
        # Add line break
        contact_paragraph.add_run("\n")
        
        # Email and phone
        contact_paragraph.add_run("npferracuti@gmail.com\n")
        contact_paragraph.add_run("416-358-4749\n")
        contact_paragraph.add_run("Toronto, Canada\n")
        
        # LinkedIn with hyperlink
        linkedin_run = contact_paragraph.add_run()
        linkedin_run.text = "Linkedin.com/in/nick-ferracuti/"
        linkedin_run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color for hyperlink
        linkedin_run.font.underline = True
        # Add hyperlink relationship
        rel_id = doc.part.relate_to('https://www.linkedin.com/in/nick-ferracuti/', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), rel_id)
        hyperlink.append(linkedin_run._r)
        contact_paragraph._p.append(hyperlink)
        
        # Add line break
        contact_paragraph.add_run("\n")
        
        # GitHub with hyperlink
        github_run = contact_paragraph.add_run()
        github_run.text = "Github.com/NFerracuti"
        github_run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color for hyperlink
        github_run.font.underline = True
        # Add hyperlink relationship
        rel_id = doc.part.relate_to('https://github.com/NFerracuti', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), rel_id)
        hyperlink.append(github_run._r)
        contact_paragraph._p.append(hyperlink)
        
        # Add spacing after contact info
        spacer = left_cell.add_paragraph()
        spacer.paragraph_format.space_before = Pt(6)
        spacer.paragraph_format.space_after = Pt(6)
        
        # Summary section
        if resume_data.get("summary"):
            summary_header = add_section_header(left_cell, "PROFESSIONAL SUMMARY")
            summary_paragraph = left_cell.add_paragraph()
            summary_paragraph.add_run(resume_data["summary"])
        
        # Skills Section
        skills_header = add_section_header(left_cell, "SKILLS")
        
        skills_data = resume_data.get("skills", "")
        if isinstance(skills_data, str) and skills_data.strip():
            # Handle structured string format from OpenAI
            sections = skills_data.split('\n\n')
            for section in sections:
                section = section.strip()
                if section:
                    if ':' in section:
                        # Category with skills
                        category, skills = section.split(':', 1)
                        category_para = left_cell.add_paragraph()
                        category_run = category_para.add_run(category.strip())
                        category_run.bold = True
                        category_run.font.size = Pt(10)  # Ensure category is 10pt
                        
                        skills_para = left_cell.add_paragraph()
                        skills_run = skills_para.add_run(skills.strip())
                        skills_run.font.size = Pt(10)  # Ensure skills are 10pt
                        skills_para.paragraph_format.space_after = Pt(6)
                    else:
                        # Standalone skill (Git or Scrum)
                        standalone_para = left_cell.add_paragraph()
                        standalone_run = standalone_para.add_run(section.strip())
                        standalone_run.bold = True
                        standalone_run.font.size = Pt(10)  # Ensure standalone skills are 10pt
                        standalone_para.paragraph_format.space_after = Pt(6)
        elif isinstance(skills_data, dict):
            # Handle dictionary format from resume_data.json
            # First add categorized skills
            for category, skills in skills_data.items():
                if category.lower() not in ['git', 'scrum']:
                    if skills:
                        category_para = left_cell.add_paragraph()
                        category_para.add_run(category.replace('_', ' ').title()).bold = True
                        skills_para = left_cell.add_paragraph()
                        skills_para.add_run(', '.join(skills))
                        skills_para.paragraph_format.space_after = Pt(6)
            
            # Then add standalone skills
            for category in ['git', 'scrum']:
                if category in skills_data and skills_data[category]:
                    standalone_para = left_cell.add_paragraph()
                    standalone_para.add_run(category.title()).bold = True
                    standalone_para.paragraph_format.space_after = Pt(6)
        else:
            # Log warning if no skills found
            self.logger.warning("No skills data found or in unexpected format")
        
        # Education Section
        if resume_data.get("education"):
            edu_header = add_section_header(left_cell, "EDUCATION")
            
            for edu in resume_data["education"]:
                # Degree/Program
                degree_paragraph = left_cell.add_paragraph()
                if edu.get("degree"):
                    degree_paragraph.add_run(edu["degree"]).bold = True
                
                # Institution
                if edu.get("institution"):
                    inst_paragraph = left_cell.add_paragraph()
                    inst_paragraph.add_run(edu["institution"])
                
                # Date and Location
                date_loc = []
                if edu.get("dates"):
                    date_loc.append(edu["dates"])
                if edu.get("location"):
                    date_loc.append(edu["location"])
                
                if date_loc:
                    loc_paragraph = left_cell.add_paragraph()
                    loc_paragraph.add_run(" | ".join(date_loc))
                
                # Add small space between education entries
                if edu != resume_data["education"][-1]:
                    spacer = left_cell.add_paragraph()
                    spacer.paragraph_format.space_before = Pt(6)
        
        # Experience section
        if resume_data.get("experience"):
            # Create header paragraph directly
            header = right_cell.add_paragraph()
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = header.add_run("PROFESSIONAL EXPERIENCE")
            run.bold = True
            # Add light grey background
            pPr = header._p.get_or_add_pPr()
            pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>'))
            header.paragraph_format.space_before = Pt(0)
            header.paragraph_format.space_after = Pt(0)
            
            # Add white spacing below header only
            spacer_after = right_cell.add_paragraph()
            spacer_after.paragraph_format.space_before = Pt(0)
            spacer_after.paragraph_format.space_after = Pt(0)
            
            for exp in resume_data["experience"]:
                # Company and title
                exp_title = right_cell.add_paragraph()
                company = exp.get('company', '')
                title = exp.get('title', '')
                if company and title:
                    exp_title.add_run(f"{title} - {company}").bold = True
                elif title:
                    exp_title.add_run(title).bold = True
                elif company:
                    exp_title.add_run(company).bold = True
                
                # Dates and location
                if exp.get('dates'):
                    date_loc = right_cell.add_paragraph()
                    date_text = []
                    date_text.append(exp['dates'])
                    if exp.get('location'):
                        date_text.append(exp['location'])
                    date_loc.add_run(" | ".join(date_text))
                
                # Description points - each gets its own bullet
                if exp.get('description'):
                    description_points = exp['description']
                    if isinstance(description_points, list):
                        for point in description_points:
                            bullet_para = right_cell.add_paragraph(style='List Bullet')
                            bullet_run = bullet_para.add_run(point)
                            bullet_run.font.size = Pt(10)  # Set description points to 10pt
                            bullet_para.paragraph_format.space_after = Pt(0)
                            bullet_para.paragraph_format.space_before = Pt(0)
                
                # Add small space between experiences
                if exp != resume_data["experience"][-1]:
                    spacer = right_cell.add_paragraph()
                    spacer.paragraph_format.space_before = Pt(6)
        
        return doc
    
    def _create_template_resume(self, resume_data):
        """Create a resume document using the template"""
        try:
            doc = DocxTemplate(self.template_path)
            
            # Set narrow margins if not using template
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.3)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)
            
            context = resume_data.copy()
            
            # Format skills as comma-separated string if it's a list
            if isinstance(context.get("skills"), list):
                context["skills_text"] = ", ".join(context["skills"])
            
            # Add formatting helpers
            context["now"] = datetime.now().strftime("%B %d, %Y")
            
            # Render the template
            doc.render(context)
            return doc
            
        except Exception as e:
            logger.error(f"Error creating resume from template: {str(e)}")
            # Fallback to basic resume
            return self._create_basic_resume(resume_data)
    
    def create_resume(self, resume_data, metadata=None):
        """Create a resume document with explicit skills validation"""
        # Make sure skills exist before proceeding
        if 'skills' not in resume_data or not resume_data['skills']:
            self.logger.warning("Skills missing from resume data! Adding default skills.")
            resume_data['skills'] = "Python, Javascript, HTML, CSS, React, Node.js"
        
        self.logger.info(f"Creating resume with skills: {resume_data['skills']}")
        
        try:
            # Create the document
            doc = self._create_basic_resume(resume_data)
            
            # Generate filename
            filename = self._generate_filename(metadata)
            output_path = os.path.join('generated_resumes', filename)
            
            # Save the document
            doc.save(output_path)
            self.logger.info(f"Resume document saved to {output_path}")
            
            # Upload to Google Drive if enabled
            if self.drive_service:
                try:
                    self._upload_to_drive(output_path)
                except Exception as e:
                    self.logger.error(f"Failed to upload to Google Drive: {str(e)}")
                    # Continue even if upload fails
            
            return output_path
            
        except Exception as e:
            self.logger.error(f"Error generating resume file: {str(e)}")
            raise
    
    def _generate_filename(self, metadata=None):
        """Generate a filename for the resume"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        if metadata and isinstance(metadata, dict):
            parts = []
            if metadata.get('job_title'):
                parts.append(metadata['job_title'].replace(' ', '_'))
            if metadata.get('company'):
                parts.append(metadata['company'].replace(' ', '_'))
            if metadata.get('job_id'):
                parts.append(metadata['job_id'])
            
            if parts:
                base_name = '_'.join(parts)
                return f"Resume_{base_name}_{timestamp}.docx"
        
        return f"Resume_{timestamp}.docx"
    
    def _upload_to_drive(self, file_path):
        """Upload file to Google Drive and return file ID"""
        if not self.drive_service:
            self.logger.warning("Google Drive service not initialized")
            return None
        
        try:
            file_metadata = {
                'name': os.path.basename(file_path),
                'mimeType': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            }
            
            media = MediaFileUpload(
                file_path,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                resumable=True
            )
            
            file = self.drive_service.files().create(
                body=file_metadata,
                media_body=media,
                fields='id'
            ).execute()
            
            file_id = file.get('id')
            if file_id:
                self.logger.info(f"File uploaded to Google Drive with ID: {file_id}")
                return file_id
            else:
                self.logger.error("File uploaded but no ID returned")
                return None
            
        except Exception as e:
            self.logger.error(f"Failed to upload file to Google Drive: {str(e)}")
            return None

    def _init_drive_service(self, credentials_path='google-service-account.json'):
        """Initialize Google Drive service using service account"""
        try:
            from google.oauth2 import service_account
            
            if not os.path.exists(credentials_path):
                self.logger.error(f"Service account credentials not found at {credentials_path}")
                return None
            
            # Load service account credentials
            credentials = service_account.Credentials.from_service_account_file(
                credentials_path,
                scopes=['https://www.googleapis.com/auth/drive.file']
            )
            
            # Build and return the service
            service = build('drive', 'v3', credentials=credentials)
            self.logger.info("Google Drive service initialized with service account")
            return service
            
        except Exception as e:
            self.logger.error(f"Failed to initialize Google Drive service: {str(e)}")
            return None

    def upload_to_google_drive(self, file_path):
        """Public method to upload file to Google Drive"""
        try:
            file_id = self._upload_to_drive(file_path)
            if file_id:
                # Create sharing permission for your email
                permission = {
                    'type': 'user',
                    'role': 'writer',
                    'emailAddress': 'npferracuti@gmail.com'
                }
                self.drive_service.permissions().create(
                    fileId=file_id,
                    body=permission,
                    sendNotificationEmail=False
                ).execute()

                # Also make it viewable by anyone with the link
                anyone_permission = {
                    'type': 'anyone',
                    'role': 'reader'
                }
                self.drive_service.permissions().create(
                    fileId=file_id,
                    body=anyone_permission
                ).execute()

                # Get sharing link
                file = self.drive_service.files().get(
                    fileId=file_id,
                    fields='webViewLink'
                ).execute()
                return file.get('webViewLink')
            return None
        except Exception as e:
            self.logger.error(f"Failed to upload to Google Drive: {str(e)}")
            return None

    def _add_experience_section(self, cell, experiences):
        """Add experience section with smaller font for descriptions"""
        self.add_section_header(cell, "EXPERIENCE")
        
        for i, exp in enumerate(experiences):
            # Company and role
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            
            # Company name
            company_run = p.add_run(f"{exp.get('company', '')}")
            company_run.bold = True
            company_run.font.size = Pt(11)
            
            p.add_run(" | ")
            
            # Job title
            title_run = p.add_run(f"{exp.get('title', '')}")
            title_run.italic = True
            title_run.font.size = Pt(11)
            
            # Date line
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            
            date_run = p.add_run(f"{exp.get('startDate', '')} - {exp.get('endDate', '')}")
            date_run.italic = True
            date_run.font.size = Pt(10)
            
            # Description bullets with SMALLER font
            description_key = "description"
            if description_key in exp and exp[description_key]:
                self.logger.info(f"Adding {len(exp[description_key])} description bullets")
                for desc in exp[description_key]:
                    p = cell.add_paragraph()
                    p.style = 'List Bullet'
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.left_indent = Inches(0.25)
                    
                    # IMPORTANT: Set font size to 9pt and add a marker for debugging
                    desc_run = p.add_run(desc)
                    desc_run.font.size = Pt(9)  # Set to size 9
                    desc_run.font.name = "Zilla Slab Medium"
                    
                    self.logger.info(f"Added bullet with font size 9pt: {desc[:30]}...")
            
            # Add spacing between experiences
            if i < len(experiences) - 1:
                p = cell.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(6)

    def _add_personal_info_section(self, cell, personal_info):
        """Add personal information section with hardcoded access to resume data"""
        # Load resume data if not already loaded
        if not hasattr(self, 'resume_data'):
            try:
                with open('assets/resume_data.json', 'r') as f:
                    self.resume_data = json.load(f)
            except Exception as e:
                self.logger.error(f"Error loading resume data: {str(e)}")
                self.resume_data = {}
        
        # Extract name - prioritize personal_info from resume_data
        name = ""
        if 'personal_info' in self.resume_data and 'name' in self.resume_data['personal_info']:
            name = self.resume_data['personal_info']['name']
        elif isinstance(personal_info, dict) and 'name' in personal_info:
            name = personal_info['name']
        
        # Add name paragraph
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        name_run = p.add_run(name)
        name_run.bold = True
        name_run.font.size = Pt(14)
        name_run.font.name = "Zilla Slab Medium"
        
        # Extract contact details - prioritize personal_info from resume_data
        email = ""
        phone = ""
        location = ""
        
        if 'personal_info' in self.resume_data:
            email = self.resume_data['personal_info'].get('email', '')
            phone = self.resume_data['personal_info'].get('phone', '')
            location = self.resume_data['personal_info'].get('location', '')
        elif isinstance(personal_info, dict):
            email = personal_info.get('email', '')
            phone = personal_info.get('phone', '')
            location = personal_info.get('location', '')
        
        # Add contact paragraph
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(0)
        
        contact_run = p.add_run(f"{email} | {phone} | {location}")
        contact_run.font.size = Pt(10)
        contact_run.font.name = "Zilla Slab Medium"
        
        # Add a spacer after the personal info
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

    def _add_skills_section(self, cell, skills):
        """Add skills section with hardcoded access to resume data"""
        self._add_section_header(cell, "SKILLS")
        
        # Load resume data if not already loaded
        if not hasattr(self, 'resume_data'):
            try:
                with open('assets/resume_data.json', 'r') as f:
                    self.resume_data = json.load(f)
                    self.logger.info(f"Loaded resume data for skills section")
            except Exception as e:
                self.logger.error(f"Error loading resume data: {str(e)}")
                self.resume_data = {}
        
        # First try to use the skills parameter
        skills_text = ""
        
        # If skills is a valid input, use it
        if isinstance(skills, list):
            skills_text = ", ".join(skills)
            self.logger.info(f"Using skills list from parameter: {skills_text[:50]}...")
        elif isinstance(skills, str) and skills.strip():
            skills_text = skills
            self.logger.info(f"Using skills string from parameter: {skills_text[:50]}...")
        # If not, extract from resume_data.json
        else:
            try:
                # Check if skills is a nested dictionary in resume_data
                if 'skills' in self.resume_data and isinstance(self.resume_data['skills'], dict):
                    all_skills = []
                    for category, skill_list in self.resume_data['skills'].items():
                        if isinstance(skill_list, list):
                            all_skills.extend(skill_list)
                    
                    skills_text = ", ".join(all_skills)
                    self.logger.info(f"Extracted skills from resume_data dictionary: {skills_text[:50]}...")
                # Check if skills is a direct list in resume_data
                elif 'skills' in self.resume_data and isinstance(self.resume_data['skills'], list):
                    skills_text = ", ".join(self.resume_data['skills'])
                    self.logger.info(f"Extracted skills from resume_data list: {skills_text[:50]}...")
            except Exception as e:
                self.logger.error(f"Error extracting skills from resume_data: {str(e)}")
        
        # Fallback if still no skills
        if not skills_text or skills_text.strip() == "":
            skills_text = "Python, JavaScript, HTML, CSS, React, Node.js"
            self.logger.warning(f"No skills found, using default: {skills_text}")
        
        # Add the skills paragraph
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        
        run = p.add_run(skills_text)
        run.font.size = Pt(10)
        
        self.logger.info(f"Added skills to document: {skills_text[:50]}...")

    def generate_resume_file(self, tailored_resume):
        """Bridge method to match what main.py is expecting"""
        self.logger.info(f"generate_resume_file called with tailored_resume containing keys: {tailored_resume.keys()}")
        
        # Convert from the OpenAI generator output format to our document format
        resume_data = {
            # Copy existing resume data from the instance
            "name": self.resume_data.get("name", ""),
            "email": self.resume_data.get("email", ""),
            "phone": self.resume_data.get("phone", ""),
            "location": self.resume_data.get("location", ""),
            
            # Add the tailored content
            "summary": tailored_resume.get("summary", ""),
            "skills": tailored_resume.get("skills", ""),
            
            # Copy other sections
            "experience": self.resume_data.get("experience", []),
            "education": self.resume_data.get("education", [])
        }
        
        self.logger.info(f"Creating resume with skills: '{resume_data['skills']}'")
        
        # Call our existing method to create the resume
        return self.create_resume(resume_data) 